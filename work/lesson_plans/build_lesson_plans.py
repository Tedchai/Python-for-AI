from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\game0\OneDrive\Desktop\Python for AI")
OUT = ROOT

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(16, 33, 63)
MUTED = RGBColor(81, 98, 122)
GOLD = RGBColor(243, 164, 0)
LIGHT_BLUE_HEX = "E8EEF5"
LIGHT_GRAY_HEX = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_text(paragraph, text, size=11, color=INK, bold=False, italic=False):
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def configure_doc(doc, running_title):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_text(header, running_title, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_text(footer, "MeritPoint Academy / 恒点学术", size=9, color=MUTED)


def add_title_block(doc, title, subtitle, meta_rows):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_font(run, size=24, color=INK, bold=True)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    set_font(sub.add_run(subtitle), size=12.5, color=MUTED, bold=True)

    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1800, 7560])
    for r, (label, value) in enumerate(meta_rows):
        cells = table.rows[r].cells
        set_cell_shading(cells[0], LIGHT_BLUE_HEX)
        paragraph_text(cells[0].paragraphs[0], label, size=10.5, color=DARK_BLUE, bold=True)
        paragraph_text(cells[1].paragraphs[0], value, size=10.5, color=INK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        paragraph_text(p, item, size=11, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        paragraph_text(p, item, size=11, color=INK)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code)
    set_font(run, size=9.5, color=RGBColor(20, 35, 55), name="Consolas")
    p.style = doc.styles["Normal"]
    p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p._p.pPr.append(shd)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE_HEX)
        paragraph_text(cell.paragraphs[0], header, size=10, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            paragraph_text(cells[i].paragraphs[0], value, size=9.5 if len(value) > 55 else 10, color=INK)
    set_table_width(table, widths)
    return table


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FFF8E8")
    p = cell.paragraphs[0]
    paragraph_text(p, f"{label}: {text}", size=10.5, color=INK, bold=True)


LESSONS = [
    {
        "file": "教案01 - Python Basics and Google Colab.docx",
        "title": "教案 01：课程导入与 Python 第一个 Notebook",
        "subtitle": "Python for AI Research | Class 01",
        "meta": [
            ("课时主题", "课程导入、Google Classroom、Colab 环境、print、变量、数据类型、注释、input、f-string"),
            ("建议时长", "60 分钟"),
            ("对应材料", "Python for AI Research - Class 01 polished expanded.pptx"),
            ("当堂产出", "学生完成 Class01_FirstNotebook，并运行 mini profile 小程序"),
            ("课后作业", "完成 8-10 个变量与输出练习，提交 Google Colab notebook 链接"),
        ],
        "objectives": [
            "学生能说出本课程 15 次课的学习路径、评分结构和作业提交方式。",
            "学生能加入 Google Classroom，并知道作业和课件分别在 Assignments 与 Class Materials 中查看。",
            "学生能在 Google Colab 中新建、重命名、运行 notebook。",
            "学生能使用 print、变量、str/int/float/bool、注释、input 和 f-string 完成简单输出。",
            "学生能解释自己代码的输出，而不是只复制运行结果。",
        ],
        "prep": [
            "教师提前打开 Class 01 课件与 Google Classroom 页面，准备班级代码 kday73zp。",
            "确认学生能登录 Google 账号并访问 colab.research.google.com。",
            "准备一个示范 notebook，文件名建议为 Class01_FirstNotebook_Demo。",
            "若学生电脑环境不一致，优先使用 Colab，避免本地 Python 安装问题消耗课堂时间。",
        ],
        "flow": [
            ("0-8 分钟", "课程开场", "讲师自我介绍，说明课程会从 Python 走到 AI 科研项目。", "学生了解课程目标与学习方式。", "学生知道课程不是单纯语法课。"),
            ("8-15 分钟", "Syllabus 与评分", "讲解 5 个阶段、每节课可见产出、35/20/25/20 评分结构与 AI 使用规范。", "学生记录作业与评分重点。", "学生能说出本课和最终项目要求。"),
            ("15-22 分钟", "加入 Classroom", "演示 classroom.google.com -> Join class -> 输入 kday73zp。", "学生现场加入课程。", "所有学生进入 The AI Toolbox。"),
            ("22-32 分钟", "Colab 入门", "演示新建 notebook、重命名、运行第一个 code cell。", "学生跟做并运行 Hello AI Research。", "Class01_FirstNotebook 文件创建完成。"),
            ("32-47 分钟", "Python 基础小步练习", "按 Read / Predict / Run / Explain 拆解 print、变量、类型、注释、input、f-string。", "学生先预测输出，再运行验证。", "学生能解释至少 3 段代码。"),
            ("47-55 分钟", "Mini profile 综合练习", "带学生完成一个输入姓名、年龄、研究兴趣、学习时间的小程序。", "学生运行并修改成自己的版本。", "一个可运行 mini profile。"),
            ("55-60 分钟", "提交与收束", "说明作业提交位置、提交前 5 项检查和下节课 Control Flow 预告。", "学生确认作业要求。", "作业路径和提交标准明确。"),
        ],
        "concepts": [
            ("print()", "把括号中的内容显示出来。第一课只要求能运行、能改动、能解释。"),
            ("变量", "变量像一个有名字的盒子，用来保存代码后面还会继续使用的值。"),
            ("数据类型", "str 是文字，int 是整数，float 是小数，bool 是真假。type() 可以检查类型。"),
            ("注释", "注释用于说明代码意图，不是逐字翻译每一行。"),
            ("input()", "从用户那里接收输入；默认得到 str，做数学计算前常需要 int() 或 float()。"),
        ],
        "codes": [
            ("第一段代码", '# Class 01\nprint("Hello, AI Research!")\nprint("My first Python notebook is running.")'),
            ("类型检查", 'age = 16\nscore = 88.5\nname = "Lina"\n\nprint(type(age))\nprint(type(score))\nprint(type(name))'),
            ("Mini profile", '# Class 01 mini profile\nname = input("Name: ")\nage = int(input("Age: "))\ntopic = input("Research interest: ")\nhours = float(input("Weekly study hours: "))\n\nprint(f"{name}, age {age}, wants to study {topic}.")\nprint(f"Next week goal: {hours + 1} hours.")'),
        ],
        "practice": [
            "把 print(\"AI\" + \" Research\") 改成自己的研究兴趣输出。",
            "创建 student_name、student_age、weekly_hours、project_topic、is_beginner 五个变量。",
            "用 type() 检查至少三个变量的类型。",
            "运行 mini profile 后，把输出结果改成更自然的一句话。",
        ],
        "homework": [
            "在 Google Colab 完成 8-10 个变量练习。",
            "至少包含 str、int、float、bool 三类以上数据。",
            "至少写 3-5 行 print 输出，2-3 行注释。",
            "确认 notebook 从上到下可以顺利运行后，把链接提交到 Google Classroom 的 Class 01 作业。",
        ],
        "teacher_notes": [
            "第一课不要追求语法完整性，重点是降低学生对代码的陌生感。",
            "遇到报错时请公开示范读错误信息，不要直接替学生修好。",
            "对基础弱的学生，要求能修改变量值和解释输出即可；对进阶学生，可加问为什么 input 得到的是 str。",
        ],
    },
    {
        "file": "教案02 - Python Control Flow.docx",
        "title": "教案 02：Python Control Flow",
        "subtitle": "Python for AI Research | Class 02",
        "meta": [
            ("课时主题", "条件判断、比较运算、逻辑结构、for loop、while loop、break、continue"),
            ("建议时长", "60-75 分钟"),
            ("对应材料", "Python Control Flow.ppsx"),
            ("当堂产出", "学生完成硬币兑换讨论、正负数判断、循环打印与正数输入验证函数"),
            ("课后作业", "写一个简单评分/分类程序，例如根据数值输出等级或风险"),
        ],
        "objectives": [
            "学生能解释 if 条件为 True/False 时程序如何选择执行路径。",
            "学生能正确使用 if、else、elif 完成二分支和多分支判断。",
            "学生能用 for 和 range 生成数字序列，并理解 start、stop、step。",
            "学生能用 while 写重复询问逻辑，并知道如何避免无限循环。",
            "学生能区分 break 与 continue 的作用，并在练习中使用其中一种。",
        ],
        "prep": [
            "打开 Python Control Flow.ppsx，并准备 Colab 示例 notebook。",
            "复习第 1 课的 // 和 %，用于 87 cents 硬币兑换讨论。",
            "准备几个输入值：87、5、-2、0、-10 到 10、正数输入验证。",
            "提醒学生本课核心不是背语法，而是预测程序会走哪条路径。",
        ],
        "flow": [
            ("0-8 分钟", "作业回顾：硬币兑换", "用 87 cents 引出 // 与 %，讨论 quarters、dimes、pennies 的计算顺序。", "学生预测每一步余数。", "学生理解整除和取余在真实任务中的用法。"),
            ("8-20 分钟", "if 语句", "讲解 if condition: 与缩进，演示 positive number 例子。", "学生输入 5、-2 预测输出。", "学生能解释缩进决定代码块。"),
            ("20-30 分钟", "if / else / elif", "从二分支到多分支，演示 positive / negative / zero。", "学生改写判断条件。", "完成三分类数字判断。"),
            ("30-42 分钟", "for loop 与 range", "拆解 for variable in sequence，讲 range(start, stop, step)。", "学生打印 -10 到 10，再只打印偶数。", "学生能用 for loop 减少重复代码。"),
            ("42-53 分钟", "while loop", "演示 countdown 与正数输入验证，强调循环终止条件。", "学生观察同一行缩进改变输出。", "学生理解 while 适合“不知道重复几次”的场景。"),
            ("53-65 分钟", "break / continue", "演示 break 提前退出、continue 跳过当前轮。", "学生判断输出序列。", "学生能说出二者差异。"),
            ("65-75 分钟", "综合练习与作业", "指导学生写评分/分类程序，并说明作业要求。", "学生选择一个分类规则。", "作业方向明确。"),
        ],
        "concepts": [
            ("Boolean expression", "条件表达式的结果只有 True 或 False，决定 if 代码块是否执行。"),
            ("Indentation", "Python 用缩进定义代码块；缩进不同，程序逻辑就不同。"),
            ("if / elif / else", "用于在多个可能路径中选择一个路径。"),
            ("for loop", "适合遍历已知序列，例如 range(-10, 11)。"),
            ("while loop", "适合重复直到某个条件不再成立。"),
            ("break / continue", "break 结束整个循环；continue 跳过当前轮剩余代码，进入下一轮。"),
        ],
        "codes": [
            ("if 基础", 'number = int(input("Enter a number: "))\n\nif number > 0:\n    print(f"{number} is a positive number.")\n\nprint("A statement outside the if statement.")'),
            ("三分类", 'number = int(input("Enter a number: "))\n\nif number > 0:\n    print("Positive number")\nelif number < 0:\n    print("Negative number")\nelse:\n    print("Zero")'),
            ("for loop 偶数", 'for i in range(-10, 11):\n    if i % 2 == 0:\n        print(f"number: {i}")'),
            ("while 输入验证", 'def take_user_input() -> int:\n    while True:\n        number = int(input("Enter a positive number: "))\n        if number > 0:\n            return number\n        print("Invalid input. Please enter a positive number.")\n\nnumber = take_user_input()\nprint(f"You entered: {number}")'),
        ],
        "practice": [
            "输入一个分数，输出 pass 或 fail。",
            "输入一个数，判断 positive、negative 或 zero。",
            "用 range(-10, 11) 打印所有数字，再只打印偶数。",
            "写一个 while 循环，直到用户输入正数才停止。",
            "用 continue 跳过偶数，只打印 0-9 中的奇数。",
        ],
        "homework": [
            "写一个简单评分/分类程序，可以是成绩等级、风险等级、温度分类或学习时间建议。",
            "必须包含至少一个 if / elif / else。",
            "至少测试 3 个不同输入，并在 notebook 中保留输出结果。",
            "写 3-5 句话解释你的判断规则，以及哪些边界输入容易出错。",
        ],
        "teacher_notes": [
            "第 2 课建议继续使用“先预测，再运行”的节奏，因为控制流最容易出现“看懂但写错”。",
            "对缩进错误不要急着纠正答案，先让学生指出哪几行属于 if 或 while 的代码块。",
            "如果时间不足，ternary operator 和 compact if 可以作为拓展，不必作为必会内容。",
        ],
    },
    {
        "file": "教案03 - Python Functions.docx",
        "title": "教案 03：Python Functions",
        "subtitle": "Python for AI Research | Class 03",
        "meta": [
            ("课时主题", "函数定义、函数调用、参数、return、参数传递、type hints、作用域、docstring"),
            ("建议时长", "60-75 分钟"),
            ("对应材料", "Python Functions.ppsx"),
            ("当堂产出", "学生把重复逻辑封装成函数，并完成 add_numbers 与 calculate_price 练习"),
            ("课后作业", "改写课堂代码，新增一个函数；若需对齐大纲，可加一个简单文本输出文件"),
        ],
        "objectives": [
            "学生能说明为什么要使用函数：减少重复、让逻辑命名、方便测试。",
            "学生能写出 def function_name(parameters): 的基本结构并调用函数。",
            "学生能区分 parameter 与 argument、positional argument 与 keyword argument。",
            "学生能使用 return 把结果传回调用者，并理解 return 后面的代码不会执行。",
            "学生能解释 local scope 与 global scope 的基本差异。",
            "学生能写一个简短 docstring 说明函数用途、输入和返回值。",
        ],
        "prep": [
            "打开 Python Functions.ppsx，并准备一个 Colab notebook。",
            "复用第 2 课 take_user_input 作为“函数为什么有用”的桥梁。",
            "准备 greet、add_numbers、calculate_price 三组示例。",
            "如需保留原大纲 Functions & File I/O，可在课末加 8 分钟 open/write 简单预告。",
        ],
        "flow": [
            ("0-8 分钟", "从重复代码引入函数", "展示重复 print 或重复计算，提问如何让代码更短、更好检查。", "学生指出重复部分。", "学生理解函数的动机。"),
            ("8-18 分钟", "函数语法与调用", "逐步拆解 def、函数名、括号、参数、冒号、函数体。演示 greet()。", "学生跟写并调用 greet。", "第一个自定义函数运行成功。"),
            ("18-30 分钟", "参数与 arguments", "用 greet(name) 和 add_numbers(num_1, num_2) 说明传入值。", "学生替换不同参数测试输出。", "学生能区分定义时参数和调用时实参。"),
            ("30-42 分钟", "return", "演示 add_numbers 只调用不 print 时没有屏幕输出；再把返回值存入 result。", "学生预测输出并改写 print。", "学生理解 return 是“把结果交回去”。"),
            ("42-55 分钟", "参数顺序与 keyword arguments", "用 calculate_price 展示顺序传错可能仍运行但结果错，或直接 TypeError。", "学生用 keyword arguments 修正顺序问题。", "学生知道关键字参数提升可读性。"),
            ("55-65 分钟", "作用域与 docstring", "演示函数内部变量在外部不可直接访问，补充 docstring 写法。", "学生解释 NameError 原因。", "学生能写简短说明文档。"),
            ("65-75 分钟", "综合练习与 File I/O 预告", "让学生写 calculate_study_goal 或 calculate_price 函数；时间允许则写入简单 txt。", "学生完成一个可复用函数。", "函数 notebook 可提交。"),
        ],
        "concepts": [
            ("Function", "一组被命名的语句，用来完成一个具体任务。"),
            ("Calling a function", "函数定义后不会自动运行，需要 function_name(arguments) 调用。"),
            ("Parameters vs arguments", "parameters 是函数定义中的占位名；arguments 是调用时传入的具体值。"),
            ("return", "把计算结果传回调用处；return 后面的同级代码不会继续执行。"),
            ("Type hints", "类型提示帮助读代码，但 Python 默认不会在运行时强制检查。"),
            ("Scope", "函数内部定义的变量通常只能在函数内部使用。"),
            ("Docstring", "函数开头的三引号说明，用来记录函数做什么、需要什么输入、返回什么结果。"),
        ],
        "codes": [
            ("greet", 'def greet(name):\n    print(f"Hello, {name}!")\n\ngreet("Sponge Bob")'),
            ("add_numbers with return", 'def add_numbers(num_1, num_2):\n    result = num_1 + num_2\n    return result\n\nresult = add_numbers(3, 5)\nprint(f"Sum: {result}")'),
            ("keyword arguments", 'def calculate_price(item: str, quantity: int, price: float) -> None:\n    total = quantity * price\n    print(f"The total price for {quantity} {item}(s) is: ${total:.2f}")\n\ncalculate_price(item="apple", quantity=3, price=0.5)\ncalculate_price(quantity=7, item="orange", price=1.2)'),
            ("optional File I/O alignment", 'summary = "Class 03 completed: functions, return, arguments, scope.\\n"\nwith open("class03_summary.txt", "w", encoding="utf-8") as file:\n    file.write(summary)'),
        ],
        "practice": [
            "写 greet(name)，输出 Hello, [name]。",
            "写 add_numbers(num_1, num_2)，返回两个数的和，并在函数外 print。",
            "写 calculate_price(item, quantity, price)，尝试 positional 与 keyword 两种调用方式。",
            "故意把参数顺序传错，观察是错误结果还是 TypeError。",
            "给一个函数补 docstring，说明 Args 和 Returns。",
        ],
        "homework": [
            "新增一个函数，例如 calculate_average、classify_score 或 calculate_study_goal。",
            "函数必须至少有 1 个参数，并至少调用 3 次。",
            "至少一次把函数返回值存入变量，再输出解释性句子。",
            "写 3-5 句话说明参数、返回值和你为什么这样封装。",
            "如教师要求对齐 File I/O：把一次结果写入 class03_result.txt。",
        ],
        "teacher_notes": [
            "第 3 课内容很多，type hints、keyword arguments、scope 可以按学生程度取舍。",
            "return 是最容易混淆的点：反复强调“return 不等于 print”。",
            "如果学生已经能写函数，可把重点转向命名、docstring 和可复用性。",
            "原课件函数部分足够充实；File I/O 建议作为 8 分钟补充或移到第 4 课与数据结构结合。",
        ],
    },
]


def build_lesson(lesson):
    doc = Document()
    configure_doc(doc, lesson["title"])
    add_title_block(doc, lesson["title"], lesson["subtitle"], lesson["meta"])

    doc.add_heading("一、教学目标", level=1)
    add_bullets(doc, lesson["objectives"])

    doc.add_heading("二、课前准备", level=1)
    add_bullets(doc, lesson["prep"])

    doc.add_heading("三、课堂流程建议", level=1)
    add_table(
        doc,
        ["时间", "环节", "教师活动", "学生活动", "阶段产出"],
        lesson["flow"],
        [1150, 1450, 2800, 2200, 1760],
    )

    doc.add_heading("四、核心概念", level=1)
    add_table(
        doc,
        ["概念", "教师讲解口径"],
        lesson["concepts"],
        [1900, 7460],
    )

    doc.add_heading("五、关键代码示例", level=1)
    for title, code in lesson["codes"]:
        doc.add_heading(title, level=2)
        add_code(doc, code)

    doc.add_heading("六、课堂练习", level=1)
    add_numbered(doc, lesson["practice"])

    doc.add_heading("七、课后作业与提交要求", level=1)
    add_bullets(doc, lesson["homework"])
    add_callout(doc, "提交提醒", "统一提交到 Google Classroom 的 Classwork -> Assignments；学生需要确认 notebook 能从上到下运行。")

    doc.add_heading("八、教师提示", level=1)
    add_bullets(doc, lesson["teacher_notes"])

    # Preset audit note embedded as hidden-ish operational footer paragraph for maintainers.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("格式说明：本文档按 compact_reference_guide 风格制作，使用 1 英寸页边距、Calibri/Microsoft YaHei、蓝色标题层级和固定宽度表格。")
    set_font(run, size=8.5, color=MUTED, italic=True)

    out_path = OUT / lesson["file"]
    doc.save(out_path)
    return out_path


def main():
    for lesson in LESSONS:
        print(build_lesson(lesson))


if __name__ == "__main__":
    main()
