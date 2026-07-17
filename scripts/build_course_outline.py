from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "course-outline.docx"

NAVY = "0C2D52"
BLUE = "1769C0"
TEAL = "0F9DB0"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
INK = "21303F"
MUTED = "5E7183"
WHITE = "FFFFFF"
GOLD = "F2A93B"


COURSE_ROWS = [
    ("1", "Kaggle、Notebook与Python入门", "项目Notebook v0"),
    ("2", "条件、循环与数据质量", "数据检查规则"),
    ("3", "函数、list与dict", "Stage 1项目卡模板"),
    ("4", "STAGE 1：Kaggle选题与项目启动", "问题、许可、数据字典、Notebook v1"),
    ("5", "pandas读取、查看与筛选", "Data Overview"),
    ("6", "数据清洗与决策日志", "clean.csv与cleaning_log.csv"),
    ("7", "可视化、简单统计与数据观察", "三张图和三条观察"),
    ("8", "STAGE 2：数据处理与EDA检查点", "清洁数据、EDA、修订后的问题"),
    ("9", "特征、目标、训练集与Baseline", "DummyRegressor与baseline MAE"),
    ("10", "一个简单模型：线性回归", "LinearRegression预测结果"),
    ("11", "MAE、错误分析与公平比较", "模型评价表和错误分析"),
    ("12", "STAGE 3：建模与结果检查点", "可复现模型证据包"),
    ("13", "从代码到Method与Results", "Method、Results、Limitations初稿"),
    ("14", "Draft Paper写作工作坊", "4–6页Draft Paper v1与PPT初稿"),
    ("15", "Stage 4：Draft Paper展示与AI论文课衔接", "Draft Paper v2、Notebook、PPT、修改清单"),
]

WHITELIST = [
    ("Student Performance (Multiple Linear Regression)", "pending manual review", "已知：作者、合成说明、10,000行、目标、许可说明；未知：CSV名/大小/缺失值", "https://www.kaggle.com/datasets/nikhil7280/student-performance-multiple-linear-regression"),
    ("Student Academic Performance Dataset", "pending manual review", "已知：250行、12列、CC0、CSV名、无缺失说明；未知：发布者显示名/精确字段/CSV大小", "https://www.kaggle.com/datasets/aryancodes12fyds/student-academic-performance-dataset"),
    ("Energy Consumption Dataset - Linear Regression", "pending manual review", "已知：作者、回归主题、CC BY 4.0；未知：来源/文件/规模/缺失值", "https://www.kaggle.com/datasets/govindaramsriram/energy-consumption-dataset-linear-regression"),
    ("Study Hours and Student Scores", "pending manual review", "已知：作者、100行、2列、随机生成、CC0；未知：CSV名/大小/缺失值", "https://www.kaggle.com/datasets/douaabennoune/study-hours-student-scores-for-linear-regression"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=90, bottom=60, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Microsoft YaHei", size=10, bold=False, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, after=4, line=1.12):
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else NAVY)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=3, line=1.15)
    for run in p.runs:
        set_run_font(run, size=10)
    if not p.runs:
        set_run_font(p.add_run(text), size=10)
    else:
        p.runs[0].text = text
    return p


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    set_run_font(title.add_run("Python for AI Research"), size=24, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    set_run_font(subtitle.add_run("Kaggle项目驱动 · 每次120分钟/35页 · 4个Stage · Draft Paper成果"), size=12, bold=True, color=TEAL)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [10656])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, after=0, line=1.05)
    set_run_font(p.add_run("课程定位：学生从教师审核的Kaggle表格数据中选题，以共同Python技术骨架完成可复现项目和4–6页论文初稿。"), size=9.5, bold=True, color=WHITE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("15节课程与项目Stage"), size=15, bold=True, color=BLUE)

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [720, 4752, 5184])
    table.style = "Table Grid"
    headers = ["课次", "课程内容", "当堂/阶段产出"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=8.5, bold=True, color=WHITE)
    for number, topic, output in COURSE_ROWS:
        row = table.add_row()
        for idx, text in enumerate((number, topic, output)):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if number in {"4", "8", "12", "15"}:
                set_cell_shading(cell, "D9EEF1")
            elif int(number) % 2 == 0:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run_font(p.add_run(text), size=7.8, bold=(number in {"4", "8", "12", "15"}), color=NAVY if number in {"4", "8", "12", "15"} else INK)
    set_table_geometry(table, [720, 4752, 5184])

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("MeritPoint Academy · Python for AI Research · Curriculum Outline"), size=8, color=MUTED)

    doc.add_page_break()
    add_heading(doc, "1. 专业课程定位")
    for text in [
        "目标人群：8–12年级高中生、Python初学者和AI科研入门学生。",
        "核心承诺：完成可信、可解释、可继续修改的研究初稿，不承诺在15节课内正式发表。",
        "技术边界：第一期统一使用结构化表格回归、DummyRegressor、LinearRegression和MAE。",
        "个性化方式：学生选择主题；教师审核数据、目标变量、模型类型和代码骨架。",
        "课程衔接：后续AI论文课继续完成文献综述、实验强化、引用、学术语言和多轮修改。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2. 课程节奏与评价")
    p = doc.add_paragraph("每节120分钟：0–10分钟目标与热身；10–30分钟概念；30–50分钟演示与追踪；50–70分钟两轮跟做；70–75分钟短休息；75–105分钟Kaggle项目工作室，其中约20分钟学生独立操作、约10分钟教师巡检/个别反馈/统一纠错；105–120分钟证据检查、提交和Exit Ticket。")
    style_paragraph(p)
    for run in p.runs:
        set_run_font(run, size=10)

    p = doc.add_paragraph("知识课使用概念—演示—追踪—两轮练习—调试—项目—Notebook整理—论文证据模板。第4、8、12、15课使用Stage检查模板，可将概念页替换为研究问题检查、学生代码解释、教师追问、修改记录和Stage Rubric；两类模板均保持33–37页及稳定的中英页面ID。")
    style_paragraph(p)
    for run in p.runs:
        set_run_font(run, size=9.5)
    rubric = doc.add_table(rows=1, cols=4)
    rubric.style = "Table Grid"
    set_table_geometry(rubric, [2664, 2664, 2664, 2664])
    for idx, text in enumerate(["35% 能运行", "30% 能解释", "25% 有证据", "10% 结构清楚"]):
        cell = rubric.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT if idx % 2 == 0 else "D9EEF1")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=9, bold=True, color=NAVY)

    add_heading(doc, "3. Kaggle候选登记表")
    p = doc.add_paragraph("以下条目尚未成为正式白名单。每期开课前必须重新核对原始来源、发布者、许可证、文件名、大小、行列、字段、缺失值、隐私风险和核验日期；未知信息不得推测。")
    style_paragraph(p)
    for run in p.runs:
        set_run_font(run, size=9.5)
    whitelist_table = doc.add_table(rows=1, cols=4)
    whitelist_table.style = "Table Grid"
    set_table_geometry(whitelist_table, [2700, 1800, 3456, 2700])
    for idx, text in enumerate(["数据集", "核验状态", "已知/未知", "Kaggle URL"]):
        cell = whitelist_table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=8.5, bold=True, color=WHITE)
    for name, status, note, url in WHITELIST:
        row = whitelist_table.add_row()
        for idx, text in enumerate((name, status, note, url)):
            cell = row.cells[idx]
            if len(whitelist_table.rows) % 2 == 0:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run_font(p.add_run(text), size=6.8, color=INK)
    set_table_geometry(whitelist_table, [2700, 1800, 3456, 2700])

    add_heading(doc, "4. 风险控制")
    for text in [
        "不接受学生临时更换为图像、文本、生物信息、化学或大型时间序列项目。",
        "不把相关性、回归系数或低MAE描述为因果证明。",
        "测试集不参与训练、特征选择、清洗决定或目标值填补；优先使用少量数值特征。",
        "核心路径固定使用test_size=0.20、random_state=42、均值DummyRegressor、LinearRegression与MAE。",
        "合成数据只能支持教学模拟结论，不能冒充真实世界研究发现。",
        "SVM、XGBoost、深度学习和复杂调参只进入Challenge，不进入全班必修。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "5. Draft Paper交付标准")
    p = doc.add_paragraph("建议4–6页、约1500–2500字。论文内容必须与最终Kaggle Notebook中的代码、数字、表格和图表一致。")
    style_paragraph(p)
    for run in p.runs:
        set_run_font(run, size=10)
    sections = [
        ("Title", "准确反映研究问题与目标变量"),
        ("Abstract", "问题、数据、方法、主要结果和限制"),
        ("Introduction", "背景、意义和研究问题"),
        ("Dataset and Variables", "Kaggle来源、许可证、访问日期、字段与目标"),
        ("Data Processing", "缺失、重复、类型和其他清洗决定"),
        ("Method", "特征、80/20拆分、random_state、Baseline、LinearRegression、MAE"),
        ("Results", "3–5张图、模型评价表和谨慎的证据陈述"),
        ("Discussion", "结果可能意味着什么，不夸大因果"),
        ("Limitations", "合成数据、样本、字段、模型和泛化边界"),
        ("Conclusion", "一句主要发现和下一步"),
        ("References", "数据、资料和外部来源"),
        ("Links", "Kaggle Dataset与Notebook链接"),
    ]
    paper_table = doc.add_table(rows=1, cols=2)
    paper_table.style = "Table Grid"
    set_table_geometry(paper_table, [2520, 8136])
    for idx, text in enumerate(["部分", "最低要求"]):
        cell = paper_table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), size=9, bold=True, color=WHITE)
    for name, requirement in sections:
        row = paper_table.add_row()
        for idx, text in enumerate((name, requirement)):
            cell = row.cells[idx]
            if len(paper_table.rows) % 2 == 0:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(text), size=8.5, bold=(idx == 0), color=NAVY if idx == 0 else INK)
    set_table_geometry(paper_table, [2520, 8136])

    add_heading(doc, "6. Stage Gate检查")
    stage_checks = [
        ("Stage 1", "问题可检验；数据和许可已审核；目标字段存在；Notebook v1可运行"),
        ("Stage 2", "原始数据保留；清洗日志完整；三张图与研究问题一致"),
        ("Stage 3", "相同拆分公平比较；MAE数字正确；错误和限制已经说明"),
        ("Stage 4", "Notebook从上到下运行；Draft Paper数字可追溯；PPT和修改清单齐全"),
    ]
    stage_table = doc.add_table(rows=1, cols=2)
    stage_table.style = "Table Grid"
    set_table_geometry(stage_table, [2160, 8496])
    for idx, text in enumerate(["检查点", "通过条件"]):
        cell = stage_table.rows[0].cells[idx]
        set_cell_shading(cell, TEAL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), size=9, bold=True, color=WHITE)
    for name, requirement in stage_checks:
        row = stage_table.add_row()
        for idx, text in enumerate((name, requirement)):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(text), size=8.5, bold=(idx == 0), color=NAVY if idx == 0 else INK)
    set_table_geometry(stage_table, [2160, 8496])

    add_heading(doc, "7. 后续AI论文课程交接")
    for text in [
        "扩展并核对文献综述、引用与数据来源；",
        "复查研究问题、特征选择和实验设计，并只在学生完全理解的前提下评估模型或实验改进；",
        "根据教师反馈进行多轮结构与语言修改，最后再评估是否具备展示、竞赛、投稿或其他发表条件。",
    ]:
        add_bullet(doc, text)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
